"""
docx_styles.py
--------------
Reusable style definitions for python-docx GDD document generation.
All style constants, color palettes, font configurations, and
paragraph format presets for a professional publisher-grade document.
"""

from typing import Tuple, Optional
try:
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Define stubs for type hints when python-docx not installed
    Pt = Inches = Cm = RGBColor = None


# ─────────────────────────────────────────────
# COLOR PALETTES
# ─────────────────────────────────────────────

class Colors:
    """All RGB color definitions for GDD documents."""

    # Document chrome
    HEADING_1 = (26, 60, 94)       # Deep navy — H1 text
    HEADING_2 = (41, 89, 133)      # Mid blue — H2 text
    HEADING_3 = (55, 111, 161)     # Light blue — H3 text
    HEADING_4 = (80, 80, 80)       # Dark grey — H4 text
    BODY_TEXT = (33, 33, 33)       # Near black — body paragraphs
    CAPTION = (100, 100, 100)      # Grey — captions and labels
    LINK = (0, 102, 204)           # Blue — hyperlinks

    # Callout boxes
    DESIGNER_NOTE_BG = (235, 244, 255)    # Light blue background
    DESIGNER_NOTE_BORDER = (41, 89, 133)  # Mid blue left border
    WARNING_BG = (255, 248, 220)          # Light amber background
    WARNING_BORDER = (255, 160, 0)        # Amber left border
    OPEN_QUESTION_BG = (255, 243, 224)    # Light orange background
    OPEN_QUESTION_BORDER = (230, 81, 0)   # Orange left border

    # Table styles
    TABLE_HEADER_BG = (26, 60, 94)        # Dark navy — table headers
    TABLE_HEADER_TEXT = (255, 255, 255)   # White — header text
    TABLE_ROW_ALT = (240, 245, 250)       # Very light blue — alternating rows
    TABLE_ROW_NORMAL = (255, 255, 255)    # White — normal rows
    TABLE_BORDER = (180, 195, 210)        # Light grey-blue — cell borders

    # Cover page
    COVER_TITLE = (26, 60, 94)            # Dark navy — game title
    COVER_SUBTITLE = (80, 80, 80)         # Grey — tagline
    COVER_META = (100, 100, 100)          # Light grey — genre/platform line
    COVER_ACCENT_LINE = (41, 89, 133)     # Blue — decorative separator

    # Status indicators
    PLACEHOLDER_BG = (245, 245, 245)      # Light grey — placeholder boxes
    PLACEHOLDER_BORDER = (180, 180, 180)  # Grey — placeholder borders
    PLACEHOLDER_TEXT = (130, 130, 130)    # Medium grey — placeholder labels

    @staticmethod
    def to_rgb_color(rgb_tuple: Tuple[int, int, int]):
        """Convert (r, g, b) tuple to RGBColor if docx is available."""
        if not DOCX_AVAILABLE:
            return None
        return RGBColor(rgb_tuple[0], rgb_tuple[1], rgb_tuple[2])


# ─────────────────────────────────────────────
# FONT CONFIGURATIONS
# ─────────────────────────────────────────────

class Fonts:
    """Font family and size definitions."""

    # Font families — universally available on Windows/Mac/Linux
    HEADING_FAMILY = "Cambria"     # Serif for headings — professional, readable
    BODY_FAMILY = "Calibri"        # Sans-serif for body — clean, modern
    CODE_FAMILY = "Courier New"    # Monospace for formulas, code, parameters
    CAPTION_FAMILY = "Calibri"     # Same as body for captions

    # Heading sizes (points)
    H1_SIZE = 24
    H2_SIZE = 18
    H3_SIZE = 14
    H4_SIZE = 12

    # Body sizes (points)
    BODY_SIZE = 11
    BODY_LARGE_SIZE = 12
    CAPTION_SIZE = 9
    CODE_SIZE = 10
    CALLOUT_SIZE = 10
    FOOTER_SIZE = 9
    HEADER_SIZE = 9

    # Cover page sizes
    COVER_TITLE_SIZE = 36
    COVER_TAGLINE_SIZE = 16
    COVER_META_SIZE = 12
    COVER_CONFIDENTIAL_SIZE = 10


# ─────────────────────────────────────────────
# PAGE LAYOUT
# ─────────────────────────────────────────────

class PageLayout:
    """Page margin and spacing constants."""

    # Page margins (inches)
    MARGIN_TOP = 1.0
    MARGIN_BOTTOM = 1.0
    MARGIN_LEFT = 1.25
    MARGIN_RIGHT = 1.25

    # Paragraph spacing (points)
    SPACE_BEFORE_H1 = 24
    SPACE_AFTER_H1 = 12
    SPACE_BEFORE_H2 = 18
    SPACE_AFTER_H2 = 8
    SPACE_BEFORE_H3 = 12
    SPACE_AFTER_H3 = 6
    SPACE_BEFORE_H4 = 8
    SPACE_AFTER_H4 = 4
    SPACE_BEFORE_BODY = 0
    SPACE_AFTER_BODY = 6
    SPACE_BEFORE_BULLET = 2
    SPACE_AFTER_BULLET = 2
    SPACE_BEFORE_CALLOUT = 10
    SPACE_AFTER_CALLOUT = 10

    # Line spacing
    BODY_LINE_SPACING = 1.15
    HEADING_LINE_SPACING = 1.0

    # Table padding (inches)
    TABLE_CELL_PADDING = 0.08

    # Callout box indent (inches)
    CALLOUT_INDENT_LEFT = 0.25
    CALLOUT_INDENT_RIGHT = 0.25


# ─────────────────────────────────────────────
# PARAGRAPH FORMAT PRESETS
# ─────────────────────────────────────────────

def apply_heading_1_style(paragraph, document) -> None:
    """Apply H1 style: large, dark navy, Cambria, page-break-before on major sections."""
    if not DOCX_AVAILABLE:
        return
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = Fonts.HEADING_FAMILY
    run.font.size = Pt(Fonts.H1_SIZE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*Colors.HEADING_1)
    paragraph.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_H1)
    paragraph.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_H1)
    paragraph.paragraph_format.keep_with_next = True


def apply_heading_2_style(paragraph) -> None:
    """Apply H2 style: medium, mid blue, Cambria."""
    if not DOCX_AVAILABLE:
        return
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = Fonts.HEADING_FAMILY
    run.font.size = Pt(Fonts.H2_SIZE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*Colors.HEADING_2)
    paragraph.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_H2)
    paragraph.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_H2)
    paragraph.paragraph_format.keep_with_next = True


def apply_heading_3_style(paragraph) -> None:
    """Apply H3 style: small heading, light blue, Cambria."""
    if not DOCX_AVAILABLE:
        return
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = Fonts.HEADING_FAMILY
    run.font.size = Pt(Fonts.H3_SIZE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*Colors.HEADING_3)
    paragraph.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_H3)
    paragraph.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_H3)
    paragraph.paragraph_format.keep_with_next = True


def apply_heading_4_style(paragraph) -> None:
    """Apply H4 style: small, dark grey, Cambria, bold."""
    if not DOCX_AVAILABLE:
        return
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = Fonts.HEADING_FAMILY
    run.font.size = Pt(Fonts.H4_SIZE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*Colors.HEADING_4)
    paragraph.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_H4)
    paragraph.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_H4)


def apply_body_style(paragraph) -> None:
    """Apply body text style: Calibri 11pt, near black."""
    if not DOCX_AVAILABLE:
        return
    for run in paragraph.runs:
        run.font.name = Fonts.BODY_FAMILY
        run.font.size = Pt(Fonts.BODY_SIZE)
        run.font.color.rgb = RGBColor(*Colors.BODY_TEXT)
    paragraph.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_BODY)
    paragraph.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_BODY)
    paragraph.paragraph_format.line_spacing = PageLayout.BODY_LINE_SPACING


def apply_code_style(paragraph) -> None:
    """Apply code/formula style: Courier New, slightly grey background."""
    if not DOCX_AVAILABLE:
        return
    for run in paragraph.runs:
        run.font.name = Fonts.CODE_FAMILY
        run.font.size = Pt(Fonts.CODE_SIZE)
        run.font.color.rgb = RGBColor(*Colors.BODY_TEXT)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def apply_caption_style(paragraph) -> None:
    """Apply caption style: Calibri 9pt, grey, italic."""
    if not DOCX_AVAILABLE:
        return
    for run in paragraph.runs:
        run.font.name = Fonts.CAPTION_FAMILY
        run.font.size = Pt(Fonts.CAPTION_SIZE)
        run.font.color.rgb = RGBColor(*Colors.CAPTION)
        run.font.italic = True
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────
# TABLE STYLES
# ─────────────────────────────────────────────

def style_table_header_row(row, document=None) -> None:
    """Style a table header row: dark navy background, white bold text."""
    if not DOCX_AVAILABLE:
        return
    for cell in row.cells:
        # Set background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*Colors.TABLE_HEADER_BG))
        tcPr.append(shd)
        # Style text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(*Colors.TABLE_HEADER_TEXT)
                run.font.name = Fonts.BODY_FAMILY
                run.font.size = Pt(Fonts.BODY_SIZE)


def style_table_data_row(row, alternate: bool = False) -> None:
    """Style a table data row with optional alternating color."""
    if not DOCX_AVAILABLE:
        return
    bg_color = Colors.TABLE_ROW_ALT if alternate else Colors.TABLE_ROW_NORMAL
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*bg_color))
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = Fonts.BODY_FAMILY
                run.font.size = Pt(Fonts.BODY_SIZE)


def set_table_borders(table) -> None:
    """Add consistent thin borders to all table cells."""
    if not DOCX_AVAILABLE:
        return
    border_hex = "{:02X}{:02X}{:02X}".format(*Colors.TABLE_BORDER)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_hex)
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ─────────────────────────────────────────────
# CALLOUT BOX HELPERS
# ─────────────────────────────────────────────

def add_designer_note(document, text: str) -> None:
    """
    Add a Designer's Note callout box to the document.
    Rendered as a bordered paragraph with light blue background.
    """
    if not DOCX_AVAILABLE:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_CALLOUT)
    p.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_CALLOUT)
    p.paragraph_format.left_indent = Inches(PageLayout.CALLOUT_INDENT_LEFT)
    p.paragraph_format.right_indent = Inches(PageLayout.CALLOUT_INDENT_RIGHT)

    # Background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*Colors.DESIGNER_NOTE_BG))
    pPr.append(shd)

    # Left border
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")  # 3pt border
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*Colors.DESIGNER_NOTE_BORDER))
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run("🎮 Designer's Note: ")
    run.font.bold = True
    run.font.name = Fonts.BODY_FAMILY
    run.font.size = Pt(Fonts.CALLOUT_SIZE)
    run.font.color.rgb = RGBColor(*Colors.HEADING_2)

    run2 = p.add_run(text)
    run2.font.name = Fonts.BODY_FAMILY
    run2.font.size = Pt(Fonts.CALLOUT_SIZE)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(*Colors.BODY_TEXT)


def add_open_question(document, text: str) -> None:
    """
    Add an Open Question callout box to the document.
    Rendered with amber/orange background for visibility.
    """
    if not DOCX_AVAILABLE:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_CALLOUT)
    p.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_CALLOUT)
    p.paragraph_format.left_indent = Inches(PageLayout.CALLOUT_INDENT_LEFT)
    p.paragraph_format.right_indent = Inches(PageLayout.CALLOUT_INDENT_RIGHT)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*Colors.OPEN_QUESTION_BG))
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*Colors.OPEN_QUESTION_BORDER))
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run("⚠ Open Question: ")
    run.font.bold = True
    run.font.name = Fonts.BODY_FAMILY
    run.font.size = Pt(Fonts.CALLOUT_SIZE)
    run.font.color.rgb = RGBColor(*Colors.OPEN_QUESTION_BORDER)

    run2 = p.add_run(text)
    run2.font.name = Fonts.BODY_FAMILY
    run2.font.size = Pt(Fonts.CALLOUT_SIZE)
    run2.font.color.rgb = RGBColor(*Colors.BODY_TEXT)


def add_placeholder_diagram(document, label: str, width_description: str = "Full width") -> None:
    """
    Add a placeholder box for a diagram or image.
    Renders as a grey bordered rectangle with centered label text.
    """
    if not DOCX_AVAILABLE:
        return
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else None

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*Colors.PLACEHOLDER_BG))
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*Colors.PLACEHOLDER_BORDER))
        pBdr.append(b)
    pPr.append(pBdr)

    run = p.add_run(f"[ DIAGRAM: {label} ]")
    run.font.name = Fonts.CODE_FAMILY
    run.font.size = Pt(Fonts.CAPTION_SIZE)
    run.font.color.rgb = RGBColor(*Colors.PLACEHOLDER_TEXT)
    run.font.bold = True

    # Add space padding above/below text within box
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)

    # Caption below the box
    caption = document.add_paragraph(f"Figure: {label}")
    apply_caption_style(caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else None


# ─────────────────────────────────────────────
# DOCUMENT SETUP HELPERS
# ─────────────────────────────────────────────

def set_document_margins(document) -> None:
    """Set page margins on all sections of the document."""
    if not DOCX_AVAILABLE:
        return
    for section in document.sections:
        section.top_margin = Inches(PageLayout.MARGIN_TOP)
        section.bottom_margin = Inches(PageLayout.MARGIN_BOTTOM)
        section.left_margin = Inches(PageLayout.MARGIN_LEFT)
        section.right_margin = Inches(PageLayout.MARGIN_RIGHT)


def add_header_footer(document, title: str, version: str, date: str) -> None:
    """
    Add header (document title) and footer (page numbers) to all sections.

    Args:
        document: python-docx Document object
        title: Game title for header
        version: Document version string (e.g., "v0.1")
        date: Document date string (e.g., "January 2025")
    """
    if not DOCX_AVAILABLE:
        return
    for section in document.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hdr_para.clear()
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hdr_para.add_run(f"{title} — Game Design Document — {version}  ")
        run.font.name = Fonts.BODY_FAMILY
        run.font.size = Pt(Fonts.HEADER_SIZE)
        run.font.color.rgb = RGBColor(*Colors.CAPTION)

        # Footer with page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        ftr_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        ftr_para.clear()
        ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "Page X of Y" using Word field codes
        run1 = ftr_para.add_run("Page ")
        run1.font.name = Fonts.BODY_FAMILY
        run1.font.size = Pt(Fonts.FOOTER_SIZE)
        run1.font.color.rgb = RGBColor(*Colors.CAPTION)

        # PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run_page = ftr_para.add_run()
        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)

        run2 = ftr_para.add_run(" of ")
        run2.font.name = Fonts.BODY_FAMILY
        run2.font.size = Pt(Fonts.FOOTER_SIZE)
        run2.font.color.rgb = RGBColor(*Colors.CAPTION)

        # NUMPAGES field
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")
        run_pages = ftr_para.add_run()
        run_pages._r.append(fldChar3)
        run_pages._r.append(instrText2)
        run_pages._r.append(fldChar4)

        run3 = ftr_para.add_run(f"   |   CONFIDENTIAL   |   {date}")
        run3.font.name = Fonts.BODY_FAMILY
        run3.font.size = Pt(Fonts.FOOTER_SIZE)
        run3.font.color.rgb = RGBColor(*Colors.CAPTION)
