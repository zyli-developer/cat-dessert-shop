"""
generate_gdd_docx.py
--------------------
Generates a professional, publisher-grade Game Design Document as a .docx file.

Usage:
    python scripts/generate_gdd_docx.py --title "My Game" --output "MyGame_GDD_v01.docx"
    python scripts/generate_gdd_docx.py --config gdd_content.json --output "MyGame_GDD_v01.docx"

Requirements:
    pip install python-docx

The script accepts either:
  (a) A JSON file containing all GDD section content (--config)
  (b) A structured Python dict passed programmatically via import
  (c) A minimal title-only run that generates a template document (--title)
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

# Ensure utils is importable when running from scripts/
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from utils.docx_styles import (
        Colors, Fonts, PageLayout,
        apply_heading_1_style, apply_heading_2_style,
        apply_heading_3_style, apply_heading_4_style,
        apply_body_style, apply_code_style, apply_caption_style,
        style_table_header_row, style_table_data_row, set_table_borders,
        add_designer_note, add_open_question, add_placeholder_diagram,
        set_document_margins, add_header_footer,
    )
    from utils.section_registry import (
        SECTIONS, SECTION_ORDER, print_section_outline,
        validate_gdd_content, validate_data_sensibility, estimate_content_size,
    )
    UTILS_AVAILABLE = True
except ImportError:
    UTILS_AVAILABLE = False
    print("Warning: utils modules not found. Running in standalone mode.")


# ─────────────────────────────────────────────
# CONTENT BUILDING
# ─────────────────────────────────────────────

def build_cover_page(doc: Any, game_data: Dict) -> None:
    """Build the cover page."""
    if not DOCX_AVAILABLE:
        return

    title = game_data.get("game_title", "UNTITLED GAME")
    tagline = game_data.get("tagline", "A new gaming experience")
    genre = game_data.get("genre", "Genre TBD")
    platform = game_data.get("platform", "Platform TBD")
    audience = game_data.get("audience", "Audience TBD")
    studio = game_data.get("studio_name", "Studio Name")
    version = game_data.get("version", "v0.1")
    date = game_data.get("date", datetime.now().strftime("%B %Y"))
    lead_designer = game_data.get("lead_designer", "Design Team")

    # Game title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.font.name = Fonts.HEADING_FAMILY
    title_run.font.size = Pt(Fonts.COVER_TITLE_SIZE)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*Colors.COVER_TITLE)
    title_para.paragraph_format.space_before = Pt(80)
    title_para.paragraph_format.space_after = Pt(8)

    # Tagline
    tag_para = doc.add_paragraph()
    tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_para.add_run(tagline)
    tag_run.font.name = Fonts.BODY_FAMILY
    tag_run.font.size = Pt(Fonts.COVER_TAGLINE_SIZE)
    tag_run.font.italic = True
    tag_run.font.color.rgb = RGBColor(*Colors.COVER_SUBTITLE)
    tag_para.paragraph_format.space_after = Pt(16)

    # Separator line (via border on paragraph)
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("─" * 40)
    sep_run.font.color.rgb = RGBColor(*Colors.COVER_ACCENT_LINE)
    sep_run.font.size = Pt(10)
    sep_para.paragraph_format.space_after = Pt(16)

    # Genre / Platform / Audience
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"{genre}  ·  {platform}  ·  {audience}"
    meta_run = meta_para.add_run(meta_text)
    meta_run.font.name = Fonts.BODY_FAMILY
    meta_run.font.size = Pt(Fonts.COVER_META_SIZE)
    meta_run.font.color.rgb = RGBColor(*Colors.COVER_META)
    meta_para.paragraph_format.space_after = Pt(60)

    # Studio name
    studio_para = doc.add_paragraph()
    studio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    studio_run = studio_para.add_run(studio)
    studio_run.font.name = Fonts.HEADING_FAMILY
    studio_run.font.size = Pt(14)
    studio_run.font.bold = True
    studio_run.font.color.rgb = RGBColor(*Colors.COVER_TITLE)
    studio_para.paragraph_format.space_after = Pt(6)

    # Document metadata
    for line in [
        f"Game Design Document",
        f"Version: {version}  ·  Date: {date}",
        f"Lead Designer(s): {lead_designer}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = Fonts.BODY_FAMILY
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(*Colors.CAPTION)
        p.paragraph_format.space_after = Pt(4)

    # Confidentiality notice
    doc.add_paragraph()
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(
        "CONFIDENTIAL — For internal use and authorized partners only. "
        "Do not distribute without written permission."
    )
    conf_run.font.name = Fonts.BODY_FAMILY
    conf_run.font.size = Pt(9)
    conf_run.font.italic = True
    conf_run.font.color.rgb = RGBColor(*Colors.CAPTION)

    # Version history table
    doc.add_page_break()
    hist_heading = doc.add_heading("Version History", level=2)
    _style_heading(hist_heading, 2)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Version", "Date", "Author", "Summary of Changes"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    if UTILS_AVAILABLE:
        style_table_header_row(table.rows[0])
        set_table_borders(table)

    # Initial row
    row = table.add_row()
    row.cells[0].text = version
    row.cells[1].text = date
    row.cells[2].text = lead_designer
    row.cells[3].text = "Initial draft"
    if UTILS_AVAILABLE:
        style_table_data_row(row, alternate=False)

    doc.add_page_break()


def _style_heading(para: Any, level: int) -> None:
    """Apply heading style to a paragraph."""
    if not DOCX_AVAILABLE or not UTILS_AVAILABLE:
        return
    if level == 1:
        apply_heading_1_style(para, None)
    elif level == 2:
        apply_heading_2_style(para)
    elif level == 3:
        apply_heading_3_style(para)
    elif level == 4:
        apply_heading_4_style(para)


def add_section(
    doc: Any,
    section_title: str,
    section_number: int,
    content: str,
    subsections: Optional[List[Dict]] = None,
    add_page_break: bool = True
) -> None:
    """
    Add a major section to the document.

    Args:
        doc: python-docx Document
        section_title: Section heading text
        section_number: Section number (1-19)
        content: Main prose content
        subsections: Optional list of {'title': str, 'content': str} dicts
        add_page_break: Whether to add a page break before this section
    """
    if not DOCX_AVAILABLE:
        return

    if add_page_break:
        doc.add_page_break()

    # Section heading
    heading_text = f"{section_number}. {section_title}"
    h1 = doc.add_heading(heading_text, level=1)
    _style_heading(h1, 1)

    # Main content
    _add_prose_content(doc, content)

    # Subsections
    if subsections:
        for sub in subsections:
            sub_title = sub.get("title", "")
            sub_content = sub.get("content", "")
            sub_subsections = sub.get("subsections", [])

            if sub_title:
                h2 = doc.add_heading(sub_title, level=2)
                _style_heading(h2, 2)

            _add_prose_content(doc, sub_content)

            for subsub in sub_subsections:
                ssub_title = subsub.get("title", "")
                ssub_content = subsub.get("content", "")
                if ssub_title:
                    h3 = doc.add_heading(ssub_title, level=3)
                    _style_heading(h3, 3)
                _add_prose_content(doc, ssub_content)


def _add_prose_content(doc: Any, content: str) -> None:
    """
    Parse and add prose content to document.
    Handles: paragraphs, bullet points (- prefix), tables (| delimited),
    code blocks (``` delimited), designer notes (> 🎮 prefix),
    open questions ([OPEN QUESTION: or [PLAYTEST: prefix), diagrams ([DIAGRAM: prefix).
    """
    if not content or not DOCX_AVAILABLE:
        return

    lines = content.split("\n")
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                if code_lines:
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph(code_text)
                    if UTILS_AVAILABLE:
                        apply_code_style(p)
                    code_lines = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table rows (| delimited)
        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows (|---|---|)
            if not all(c in "-| " for c in line.strip()):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            in_table = False
            if len(table_rows) >= 2:
                _add_table_from_rows(doc, table_rows)
            table_rows = []

        # Designer Note
        if line.strip().startswith("> 🎮") or line.strip().startswith("> Designer"):
            note_text = line.strip().lstrip("> 🎮").strip()
            if note_text.startswith("Designer's Note:"):
                note_text = note_text[len("Designer's Note:"):].strip()
            if UTILS_AVAILABLE:
                add_designer_note(doc, note_text)
            i += 1
            continue

        # Open Question / Playtest flag
        if "[OPEN QUESTION:" in line or "[PLAYTEST:" in line:
            if UTILS_AVAILABLE:
                # Extract the question text
                text = line.strip()
                add_open_question(doc, text)
            else:
                p = doc.add_paragraph(line.strip())
            i += 1
            continue

        # Diagram placeholder
        if line.strip().startswith("[DIAGRAM:"):
            diagram_label = line.strip().lstrip("[DIAGRAM:").rstrip("]").strip()
            if UTILS_AVAILABLE:
                add_placeholder_diagram(doc, diagram_label)
            i += 1
            continue

        # Headings (## prefix)
        if line.startswith("#### "):
            h = doc.add_heading(line[5:], level=4)
            _style_heading(h, 4)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            _style_heading(h, 3)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            _style_heading(h, 2)

        # Bullet points
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            bullet_text = line.strip().lstrip("-").lstrip("*").strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet_text)
            run.font.name = Fonts.BODY_FAMILY
            run.font.size = Pt(Fonts.BODY_SIZE)
            p.paragraph_format.space_before = Pt(PageLayout.SPACE_BEFORE_BULLET)
            p.paragraph_format.space_after = Pt(PageLayout.SPACE_AFTER_BULLET)

        # Numbered list
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".)" and line[2] == " ":
            list_text = line[2:].strip()
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(list_text)
            run.font.name = Fonts.BODY_FAMILY
            run.font.size = Pt(Fonts.BODY_SIZE)

        # Empty line
        elif not line.strip():
            pass  # natural paragraph break

        # Regular paragraph
        else:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                if UTILS_AVAILABLE:
                    apply_body_style(p)

        i += 1

    # Flush any open table
    if in_table and len(table_rows) >= 2:
        _add_table_from_rows(doc, table_rows)


def _add_table_from_rows(doc: Any, rows: List[List[str]]) -> None:
    """Add a formatted table to the document from row data."""
    if not rows or not DOCX_AVAILABLE:
        return

    n_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < n_cols:
                table.rows[r_idx].cells[c_idx].text = cell_text

    if UTILS_AVAILABLE:
        if len(table.rows) > 0:
            style_table_header_row(table.rows[0])
        for r_idx in range(1, len(table.rows)):
            style_table_data_row(table.rows[r_idx], alternate=(r_idx % 2 == 0))
        set_table_borders(table)

    doc.add_paragraph()  # Space after table


def add_toc_placeholder(doc: Any) -> None:
    """Add a Table of Contents placeholder (Word will update on open)."""
    if not DOCX_AVAILABLE:
        return
    h = doc.add_heading("Table of Contents", level=1)
    _style_heading(h, 1)

    # Insert Word TOC field
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    note = doc.add_paragraph(
        "[ Right-click → Update Field to populate this Table of Contents ]"
    )
    note_run = note.runs[0]
    note_run.font.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(*Colors.CAPTION)
    note.paragraph_format.space_after = Pt(12)
    doc.add_page_break()


# ─────────────────────────────────────────────
# MAIN GENERATION FUNCTION
# ─────────────────────────────────────────────

def generate_gdd_docx(
    game_data: Dict,
    output_path: str,
    include_toc: bool = True,
    include_template_sections: bool = True,
    strict: bool = False
) -> str:
    """
    Generate a complete GDD .docx file.

    Args:
        game_data: Dictionary containing game metadata and section content.
            Required keys: game_title
            Optional keys: tagline, genre, platform, audience, studio_name,
                          version, date, lead_designer, sections (dict of section content)
        output_path: Path for the output .docx file
        include_toc: Whether to include a table of contents
        include_template_sections: Whether to add template placeholder sections

    Returns:
        Absolute path to the generated file.
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required. Install with:\n"
            "  pip install python-docx"
        )

    # Pre-export validation
    sections_to_validate = game_data.get("sections", {})
    if UTILS_AVAILABLE and sections_to_validate:
        for warning in validate_gdd_content(sections_to_validate):
            print(f"  WARNING: {warning}")
        sensibility_warnings = validate_data_sensibility(
            sections_to_validate, strict=strict
        )
        for warning in sensibility_warnings:
            print(f"  WARNING: {warning}")
        if strict and sensibility_warnings:
            raise SystemExit(
                "STRICT MODE: Export aborted due to unsourced metrics or "
                "placeholders in business sections. Fix the warnings above "
                "or remove --strict to export with warnings."
            )
        size_info = estimate_content_size(sections_to_validate)
        for warning in size_info["warnings"]:
            print(f"  WARNING: {warning}")

    doc = Document()

    # Set document margins
    if UTILS_AVAILABLE:
        set_document_margins(doc)

    # Add header/footer
    game_title = game_data.get("game_title", "Untitled Game")
    version = game_data.get("version", "v0.1")
    date = game_data.get("date", datetime.now().strftime("%B %Y"))

    if UTILS_AVAILABLE:
        add_header_footer(doc, game_title, version, date)

    # Cover page
    build_cover_page(doc, game_data)

    # Table of Contents
    if include_toc:
        add_toc_placeholder(doc)

    # Sections
    sections_content = game_data.get("sections", {})

    section_defs = SECTIONS if UTILS_AVAILABLE else {}
    section_order = SECTION_ORDER if UTILS_AVAILABLE else list(sections_content.keys())

    for idx, section_key in enumerate(section_order):
        if section_key not in sections_content and include_template_sections:
            # Generate template placeholder section
            if UTILS_AVAILABLE and section_key in SECTIONS:
                sdef = SECTIONS[section_key]
                placeholder_content = _generate_placeholder_section(sdef, game_data)
                add_section(
                    doc,
                    sdef["name"],
                    sdef["order"],
                    placeholder_content,
                    add_page_break=(idx > 0)
                )
        elif section_key in sections_content:
            content = sections_content[section_key]
            if UTILS_AVAILABLE and section_key in SECTIONS:
                sdef = SECTIONS[section_key]
                add_section(
                    doc,
                    sdef["name"],
                    sdef["order"],
                    content if isinstance(content, str) else content.get("content", ""),
                    subsections=content.get("subsections") if isinstance(content, dict) else None,
                    add_page_break=(idx > 0)
                )
            else:
                add_section(doc, section_key.replace("_", " ").title(), idx + 1,
                            content if isinstance(content, str) else "",
                            add_page_break=(idx > 0))

    # Ensure output directory exists
    output_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    print(f"✓ GDD document generated: {abs_path}")
    return abs_path


def _generate_placeholder_section(section_def: Dict, game_data: Dict) -> str:
    """Generate template placeholder content for a GDD section."""
    game_title = game_data.get("game_title", "[GAME TITLE]")
    elements = section_def.get("key_elements", [])

    lines = [
        f"[This section covers the {section_def['name']} for {game_title}.]",
        "",
        f"Target length: {section_def['recommended_words']} words",
        "",
        "Required elements:",
    ]
    for elem in elements:
        lines.append(f"- {elem.replace('_', ' ').title()}: [content required]")

    if section_def.get("template_file"):
        lines.append("")
        lines.append(f"Reference: {section_def['template_file']}")

    lines.append("")
    lines.append("[OPEN QUESTION: Replace this placeholder with actual content]")

    return "\n".join(lines)


# ─────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a professional Game Design Document (.docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate a template GDD with just a title:
  python generate_gdd_docx.py --title "Echo Chamber" --output "EchoChamber_GDD_v01.docx"

  # Generate from a JSON config file:
  python generate_gdd_docx.py --config gdd_content.json --output "MyGame_GDD_v01.docx"

  # Print section outline:
  python generate_gdd_docx.py --list-sections
        """
    )
    parser.add_argument("--title", help="Game title (for quick template generation)")
    parser.add_argument("--studio", default="Studio Name", help="Studio name")
    parser.add_argument("--genre", default="Genre TBD", help="Game genre")
    parser.add_argument("--platform", default="Platform TBD", help="Target platforms")
    parser.add_argument("--audience", default="Audience TBD", help="Target audience")
    parser.add_argument("--tagline", default="", help="Game tagline")
    parser.add_argument("--version", default="v0.1", help="Document version")
    parser.add_argument("--config", help="Path to JSON config file with GDD content")
    parser.add_argument("--output", default="GDD_output.docx", help="Output file path")
    parser.add_argument("--no-toc", action="store_true", help="Skip table of contents")
    parser.add_argument("--strict", action="store_true",
                        help="Fail export if unsourced metrics or placeholders remain in business sections")
    parser.add_argument("--list-sections", action="store_true",
                        help="Print GDD section outline and exit")

    args = parser.parse_args()

    if args.list_sections:
        if UTILS_AVAILABLE:
            print(print_section_outline())
        else:
            print("Section registry not available. Ensure utils/ is in the path.")
        return

    if not DOCX_AVAILABLE:
        print("ERROR: python-docx is not installed.")
        print("Install it with: pip install python-docx")
        sys.exit(1)

    # Load from config file if provided
    if args.config:
        try:
            with open(args.config, "r", encoding="utf-8") as f:
                game_data = json.load(f)
        except FileNotFoundError:
            print(f"ERROR: Config file not found: {args.config}")
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"ERROR: Invalid JSON in config file: {e}")
            sys.exit(1)
    elif args.title:
        game_data = {
            "game_title": args.title,
            "tagline": args.tagline or f"A new {args.genre} experience",
            "genre": args.genre,
            "platform": args.platform,
            "audience": args.audience,
            "studio_name": args.studio,
            "version": args.version,
            "date": datetime.now().strftime("%B %Y"),
            "lead_designer": "Design Team",
            "sections": {}
        }
    else:
        parser.print_help()
        print("\nERROR: Provide either --title or --config")
        sys.exit(1)

    try:
        output_path = generate_gdd_docx(
            game_data=game_data,
            output_path=args.output,
            include_toc=not args.no_toc,
            include_template_sections=True,
            strict=args.strict
        )
        print(f"\n✓ Success! Open in Word and right-click the TOC to update page numbers.")
        print(f"  File: {output_path}")
    except Exception as e:
        print(f"ERROR during generation: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
